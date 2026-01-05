from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
import os

IMAGE_DIR = "histograms"
OUTPUT_DOC = "histogram_report.docx"

doc = Document()

# A4 ayarları
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1)
section.bottom_margin = Cm(1)
section.left_margin = Cm(1)
section.right_margin = Cm(1)

images = sorted([
    f for f in os.listdir(IMAGE_DIR)
    if f.lower().endswith(".png")
])

for i, img_name in enumerate(images):
    if i % 8 == 0:
        table = doc.add_table(rows=2, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            row.height = Cm(6.5)
            for cell in row.cells:
                cell.width = Cm(4.8)

    cell = table.rows[(i % 8) // 4].cells[(i % 8) % 4]
    p = cell.paragraphs[0]
    run = p.add_run()
    run.add_picture(os.path.join(IMAGE_DIR, img_name),
                    width=Cm(4.8),
                    height=Cm(6.5))

doc.save(OUTPUT_DOC)

print("Word dosyası oluşturuldu:", OUTPUT_DOC)
